#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
gitdiff2pdf - Generate PR-style PDF and Word documents from unified git diffs.

Features:
  - Unified view (default): deletions red, additions green, context grey.
  - Side-by-side view: old (left) vs. new (right).
  - File badges, hunk headers, line numbers.
  - Robust parser for diff --git / --- +++ / rename / hunks.
  - Encoding auto-detection: UTF-8, UTF-8-BOM, UTF-16 LE/BE, Latin-1.
  - System fonts (Windows: Consolas/Segoe UI, Linux: DejaVu) with safe fallbacks.
  - Pagination: keep-together per file/hunk, widow/orphan protection.
  - Optional Word (.docx) output via python-docx.
"""

from __future__ import annotations

import argparse
import datetime as dt
import os
import platform
import re
import sys
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

import fitz  # PyMuPDF

# ---------------------------------------------------------------------------
#  Colour helper
# ---------------------------------------------------------------------------

def rgb(r: int, g: int, b: int) -> Tuple[float, float, float]:
    """Convert 0-255 RGB integers to a 0.0-1.0 tuple used by PyMuPDF."""
    return (r / 255.0, g / 255.0, b / 255.0)


def _rgb_to_hex(rgb_tuple: Tuple[float, float, float]) -> str:
    """Convert a (0..1, 0..1, 0..1) tuple to a 6-digit hex string (no '#')."""
    r, g, b = rgb_tuple
    return f"{int(r * 255):02X}{int(g * 255):02X}{int(b * 255):02X}"

# ---------------------------------------------------------------------------
#  Text utilities
# ---------------------------------------------------------------------------

def text_width(s: str, fontname: str, fontsize: float) -> float:
    """Measure the rendered width of *s* in the given font."""
    return fitz.get_text_length(s, fontname=fontname, fontsize=fontsize)


# Characters that are invisible or cause layout problems.
_INVIS_CHARS = "".join([
    "\ufeff",                                                                   # BOM
    "\u200b\u200c\u200d\u2060",                                                 # zero-width
    "\u00a0\u202f\u205f",                                                       # NBSP / narrow / medium
    "\u2000\u2001\u2002\u2003\u2004\u2005\u2006\u2007\u2008\u2009\u200a",       # various spaces
])
_INVISIBLE_TRANS = dict.fromkeys(map(ord, _INVIS_CHARS), None)


def strip_invisibles(s: str) -> str:
    """Remove invisible / problematic Unicode characters."""
    return s.translate(_INVISIBLE_TRANS)


def clean_leading_artifacts(text: str) -> str:
    """Strip ellipsis / dot / bullet artefacts at the very start of a document."""
    t = text.lstrip()
    patterns = [
        "···", "...", "\u2026",                                    # ellipsis variants
        "•", "\u2022",                                             # bullets
        "‧", "\u2027", "∙", "\u2219", "⋅", "\u22c5",             # dots
    ]
    for p in patterns:
        if t.startswith(p):
            t = t[len(p):].lstrip()
            break
    return t if len(t) < len(text) else text

# ---------------------------------------------------------------------------
#  I/O helpers
# ---------------------------------------------------------------------------

def read_text(path: str) -> str:
    """Read *path* (or ``-`` for STDIN) with robust encoding fallback, then sanitise."""
    data = sys.stdin.buffer.read() if path == "-" else open(path, "rb").read()
    for enc in ("utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    return clean_leading_artifacts(strip_invisibles(text))


def norm_lines(s: str) -> List[str]:
    """Normalise line breaks and return individual lines (keeping trailing \\n)."""
    return s.replace("\r\n", "\n").replace("\r", "\n").splitlines(True)


def sanitize_path(s: str) -> str:
    """Cut at the first non-path-safe character (guards against copy/paste artefacts)."""
    allowed = set("abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789._- /\\")
    out: list[str] = []
    for ch in s:
        if ch in allowed:
            out.append(ch)
        else:
            break
    return "".join(out).strip()

# ---------------------------------------------------------------------------
#  Theme
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class Theme:
    # UI
    ui_text:     Tuple[float, float, float]
    ui_subtle:   Tuple[float, float, float]
    header_line: Tuple[float, float, float]
    # Backgrounds
    bg_added:    Tuple[float, float, float]
    bg_removed:  Tuple[float, float, float]
    bg_context:  Tuple[float, float, float]
    bg_hunk:     Tuple[float, float, float]
    # Foregrounds
    tx_added:    Tuple[float, float, float]
    tx_removed:  Tuple[float, float, float]
    tx_context:  Tuple[float, float, float]
    tx_hunk:     Tuple[float, float, float]
    # Side bars
    bar_added:   Tuple[float, float, float]
    bar_removed: Tuple[float, float, float]


LIGHT = Theme(
    ui_text=rgb(40, 40, 40),       ui_subtle=rgb(125, 125, 125),
    header_line=rgb(210, 210, 210),
    bg_added=rgb(225, 245, 234),   bg_removed=rgb(252, 232, 232),
    bg_context=rgb(247, 247, 249), bg_hunk=rgb(232, 240, 252),
    tx_added=rgb(22, 125, 57),     tx_removed=rgb(178, 36, 30),
    tx_context=rgb(60, 60, 60),
    tx_hunk=rgb(30, 90, 200),
    bar_added=rgb(34, 170, 84),    bar_removed=rgb(220, 64, 52),
)

DARK = Theme(
    ui_text=rgb(230, 230, 230),    ui_subtle=rgb(170, 170, 170),
    header_line=rgb(80, 80, 80),
    bg_added=rgb(36, 64, 52),      bg_removed=rgb(72, 40, 40),
    bg_context=rgb(36, 36, 40),    bg_hunk=rgb(44, 60, 84),
    tx_added=rgb(170, 235, 190),   tx_removed=rgb(255, 170, 170),
    tx_context=rgb(230, 230, 230),
    tx_hunk=rgb(160, 190, 255),
    bar_added=rgb(60, 200, 110),   bar_removed=rgb(240, 90, 80),
)

# ---------------------------------------------------------------------------
#  Fonts
# ---------------------------------------------------------------------------

@dataclass
class Fonts:
    ui:        str
    ui_bold:   str
    mono:      str
    mono_bold: str


def safe_font(fontname: str, fallback: str) -> str:
    """Return *fontname* if PyMuPDF can resolve it, otherwise *fallback*."""
    try:
        if not fontname or " " in fontname or "\t" in fontname:
            return fallback
        fitz.get_text_length("M", fontname=fontname, fontsize=10)
        return fontname
    except Exception:
        return fallback


def _load_font(path: Optional[str]) -> Optional[str]:
    """Try to load a TTF/OTF file; return its font name or ``None``."""
    if not path or not os.path.isfile(path):
        return None
    try:
        return fitz.Font(fontfile=path).name
    except Exception:
        return None


def detect_system_fonts() -> Fonts:
    """Detect available system fonts (Windows / Linux) with Courier as fallback."""
    ui = ui_b = "courier"
    mono = mono_b = "courier"

    system = platform.system().lower()

    if "windows" in system:
        wins = os.environ.get("WINDIR", r"C:\Windows")
        candidates = {
            "mono":   [rf"{wins}\Fonts\consola.ttf",  rf"{wins}\Fonts\cour.ttf"],
            "mono_b": [rf"{wins}\Fonts\consolab.ttf", rf"{wins}\Fonts\courbd.ttf"],
            "ui":     [rf"{wins}\Fonts\segoeui.ttf",  rf"{wins}\Fonts\arial.ttf"],
            "ui_b":   [rf"{wins}\Fonts\segoeuib.ttf", rf"{wins}\Fonts\arialbd.ttf"],
        }
        for key, paths in candidates.items():
            found = _load_font(next((p for p in paths if os.path.isfile(p)), None))
            if found:
                if key == "mono":     mono = found
                elif key == "mono_b": mono_b = found
                elif key == "ui":     ui = found
                elif key == "ui_b":   ui_b = found

    elif "linux" in system:
        base = "/usr/share/fonts/truetype/dejavu"
        mono   = _load_font(f"{base}/DejaVuSansMono.ttf")      or mono
        mono_b = _load_font(f"{base}/DejaVuSansMono-Bold.ttf")  or mono_b
        ui     = _load_font(f"{base}/DejaVuSans.ttf")           or ui
        ui_b   = _load_font(f"{base}/DejaVuSans-Bold.ttf")      or ui_b

    return Fonts(ui=ui, ui_bold=ui_b, mono=mono, mono_bold=mono_b)

# ---------------------------------------------------------------------------
#  Layout constants
# ---------------------------------------------------------------------------

@dataclass
class Layout:
    margin:            float = 44.0
    font_size:         float = 9.5
    line_gap:          float = 2.2
    hunk_gap_y:        float = 8.0
    col_gap:           float = 16.0
    gutter_gap:        float = 6.0
    gutter_chars:      int   = 5
    gap_badge_to_hunk: float = 2.0
    gap_hunk_to_code:  float = 4.0
    block_gap_y:       float = 6.0
    min_rows_on_page:  int   = 3

# ---------------------------------------------------------------------------
#  Diff data model
# ---------------------------------------------------------------------------

@dataclass
class DiffLine:
    kind:    str                    # 'ctx' | 'del' | 'add'
    text:    str
    old_num: Optional[int] = None
    new_num: Optional[int] = None


@dataclass
class Hunk:
    header:    str
    old_start: int
    old_count: int
    new_start: int
    new_count: int
    lines:     List[DiffLine] = field(default_factory=list)


@dataclass
class DiffFile:
    old_path: str        = ""
    new_path: str        = ""
    hunks:    List[Hunk] = field(default_factory=list)

# ---------------------------------------------------------------------------
#  Diff parser
# ---------------------------------------------------------------------------

_HUNK_RE = re.compile(r"@@\s*-(\d+)(?:,(\d+))?\s+\+(\d+)(?:,(\d+))?\s*@@(?:(.*))?$")


def _parse_path_from_diff_git(line: str) -> Tuple[Optional[str], Optional[str]]:
    """Extract old/new path from a ``diff --git a/… b/…`` line."""
    parts = line.strip().split()
    if len(parts) >= 4 and parts[0] == "diff" and parts[1] == "--git":
        a = parts[2][2:] if parts[2].startswith("a/") else parts[2]
        b = parts[3][2:] if parts[3].startswith("b/") else parts[3]
        return a, b
    return None, None


def _parse_path_line(line: str) -> Optional[str]:
    """Parse ``--- a/path`` or ``+++ b/path`` into a clean file path."""
    rest = line[4:].strip()
    if rest == "/dev/null":
        return None
    if rest.startswith("a/") or rest.startswith("b/"):
        rest = rest[2:]
    return strip_invisibles(sanitize_path(rest)) or None


def parse_unified_diff(text: str, tabsize: int, debug: bool = False) -> List[DiffFile]:
    """Parse a unified diff into a list of :class:`DiffFile` objects."""
    lines = norm_lines(text)

    files: List[DiffFile]        = []
    current: Optional[DiffFile]  = None
    current_hunk: Optional[Hunk] = None
    saw_any_hunk                 = False
    rename_from: Optional[str]   = None
    rename_to: Optional[str]     = None

    # -- Pass 1: build file / hunk / line structures -------------------------

    for raw in lines:
        line = strip_invisibles(raw.rstrip("\n"))

        # Skip metadata lines
        if (line.startswith("index ")
                or line.startswith("new file mode")
                or line.startswith("deleted file mode")
                or line.startswith("Binary files ")
                or line.startswith("\\ No newline at end of file")):
            continue

        # diff --git header
        if line.startswith("diff --git "):
            if current:
                files.append(current)
            a, b = _parse_path_from_diff_git(line)
            current = DiffFile(old_path=a or "", new_path=b or "")
            current_hunk = None
            rename_from = rename_to = None
            continue

        # Rename tracking
        if line.startswith("rename from "):
            rename_from = line[len("rename from "):].strip()
            continue
        if line.startswith("rename to "):
            rename_to = line[len("rename to "):].strip()
            continue

        # Old / new path
        if line.startswith("--- "):
            if current is None:
                current = DiffFile()
            p = _parse_path_line(line)
            if p is not None:
                current.old_path = p
            continue

        if line.startswith("+++ "):
            if current is None:
                current = DiffFile()
            p = _parse_path_line(line)
            if p is not None:
                current.new_path = p
            if rename_from and not current.old_path:
                current.old_path = rename_from
            if rename_to and not current.new_path:
                current.new_path = rename_to
            continue

        # Hunk header
        m = _HUNK_RE.match(line)
        if m:
            saw_any_hunk = True
            if current is None:
                current = DiffFile()

            old_start = int(m.group(1))
            old_count = int(m.group(2) or "1")
            new_start = int(m.group(3))
            new_count = int(m.group(4) or "1")

            header = f"@@ -{old_start}"
            if m.group(2):
                header += f",{old_count}"
            header += f" +{new_start}"
            if m.group(4):
                header += f",{new_count}"
            header += " @@"

            current_hunk = Hunk(
                header=header,
                old_start=old_start, old_count=old_count,
                new_start=new_start, new_count=new_count,
            )
            current.hunks.append(current_hunk)
            continue

        # Diff content lines
        if current_hunk is None:
            continue

        if line.startswith("+") and not line.startswith("+++ "):
            current_hunk.lines.append(
                DiffLine(kind="add", text=strip_invisibles(line[1:].expandtabs(tabsize))))
        elif line.startswith("-") and not line.startswith("--- "):
            current_hunk.lines.append(
                DiffLine(kind="del", text=strip_invisibles(line[1:].expandtabs(tabsize))))
        else:
            t = line[1:] if line.startswith(" ") else line
            current_hunk.lines.append(
                DiffLine(kind="ctx", text=strip_invisibles(t.expandtabs(tabsize))))

    if current:
        files.append(current)

    if not saw_any_hunk:
        return []

    # -- Pass 2: assign line numbers -----------------------------------------

    for f in files:
        if not f.old_path and not f.new_path:
            f.old_path = "(Unnamed OLD)"
            f.new_path = "(Unnamed NEW)"

        for h in f.hunks:
            old_ln = h.old_start
            new_ln = h.new_start
            for dl in h.lines:
                if dl.kind == "ctx":
                    dl.old_num = old_ln
                    dl.new_num = new_ln
                    old_ln += 1
                    new_ln += 1
                elif dl.kind == "del":
                    dl.old_num = old_ln
                    old_ln += 1
                elif dl.kind == "add":
                    dl.new_num = new_ln
                    new_ln += 1

    return files


# ---------------------------------------------------------------------------
#  Text wrapping
# ---------------------------------------------------------------------------

def _wrap_text(s: str, max_w: float, fontname: str, fontsize: float) -> List[str]:
    """Wrap *s* into lines that fit *max_w* pixels, preferring word boundaries."""
    if not s:
        return [""]
    if text_width(s, fontname, fontsize) <= max_w:
        return [s]

    out: List[str] = []
    rest = s
    while rest:
        if text_width(rest, fontname, fontsize) <= max_w:
            out.append(rest)
            break
        # Binary search for the longest fitting prefix
        lo, hi, cut = 1, len(rest), 1
        while lo <= hi:
            mid = (lo + hi) // 2
            if text_width(rest[:mid], fontname, fontsize) <= max_w:
                cut = mid
                lo = mid + 1
            else:
                hi = mid - 1
        # Prefer breaking at whitespace
        segment = rest[:cut]
        ws = max(segment.rfind(" "), segment.rfind("\t"))
        if ws >= 0 and ws >= int(0.6 * cut):
            out.append(segment[:ws].rstrip())
            rest = rest[ws + 1:]
        else:
            out.append(segment)
            rest = rest[cut:]
    return out

# ---------------------------------------------------------------------------
#  PDF Renderer
# ---------------------------------------------------------------------------

class Renderer:
    """Renders parsed diff files into a multi-page PDF using PyMuPDF."""

    def __init__(self, theme: Theme, layout: Layout, landscape: bool, fonts: Fonts):
        self.theme     = theme
        self.layout    = layout
        self.landscape = landscape
        self.doc       = fitz.open()

        self.page:   Optional[fitz.Page] = None
        self.y_base: Optional[float]     = None
        self.title:  Optional[str]       = None

        self.ui_font   = safe_font(fonts.ui,        fallback="courier")
        self.ui_bold   = safe_font(fonts.ui_bold,   fallback="courier-bold")
        self.mono_font = safe_font(fonts.mono,      fallback="courier")
        self.mono_bold = safe_font(fonts.mono_bold,  fallback="courier-bold")

    # -- Page management -----------------------------------------------------

    def new_page(self) -> fitz.Page:
        base = fitz.paper_rect("a4")
        if self.landscape:
            return self.doc.new_page(width=base.height, height=base.width)
        return self.doc.new_page(width=base.width, height=base.height)

    def box(self, page: fitz.Page) -> Tuple[float, float, float, float]:
        m = self.layout.margin
        return (m, m, page.rect.width - m, page.rect.height - m)

    def page_capacity(self) -> float:
        assert self.page is not None
        _, y0, _, y1 = self.box(self.page)
        return y1 - (y0 + 28)

    def space_left(self) -> float:
        assert self.page is not None and self.y_base is not None
        _, _, _, y1 = self.box(self.page)
        return y1 - (self.y_base - self.layout.font_size)

    def ensure_y(self, rows_h: float):
        """Start a new page if there is not enough vertical space for *rows_h*."""
        assert self.page is not None and self.y_base is not None and self.title is not None
        if self.space_left() >= rows_h:
            return
        self.page = self.new_page()
        self.draw_header(self.page, self.title)
        _, ny0, _, _ = self.box(self.page)
        self.y_base = ny0 + 28 + self.layout.font_size

    def start_if_needed(self, title: str):
        if self.page is None:
            self.page = self.new_page()
            self.title = title
            self.draw_header(self.page, title)
            _, y0, _, _ = self.box(self.page)
            self.y_base = y0 + 28 + self.layout.font_size

    def widow_check_before_hunk(self, line_h: float):
        needed = line_h + self.layout.gap_hunk_to_code + (self.layout.min_rows_on_page * line_h)
        if self.space_left() < needed:
            self.ensure_y(10_000)

    # -- Drawing primitives --------------------------------------------------

    def draw_header(self, page: fitz.Page, title: str):
        x0, y0, x1, _ = self.box(page)
        fs = self.layout.font_size
        page.insert_text((x0, y0), title,
                         fontname=self.ui_bold, fontsize=fs + 2, color=self.theme.ui_text)
        stamp = dt.datetime.now().strftime("%Y-%m-%d %H:%M")
        tw = text_width(stamp, self.ui_font, fs)
        page.insert_text((x1 - tw, y0), stamp,
                         fontname=self.ui_font, fontsize=fs, color=self.theme.ui_subtle)
        y = y0 + fs + 2 + fs + 4
        page.draw_line((x0, y - 4), (x1, y - 4), color=self.theme.header_line, width=0.8)

    def draw_footer_page_numbers(self):
        total = self.doc.page_count
        if total == 0:
            return
        fs = self.layout.font_size
        for i in range(total):
            page = self.doc.load_page(i)
            x0, _, x1, y1 = self.box(page)
            label = f"Page {i + 1} / {total}"
            tw = text_width(label, self.ui_font, fs)
            page.insert_text(((x0 + x1) / 2 - tw / 2, y1), label,
                             fontname=self.ui_font, fontsize=fs, color=self.theme.ui_subtle)

    def draw_file_badge(self, label: str):
        assert self.page is not None and self.y_base is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        pad_x = 7

        top = self.y_base - fs
        bottom = top + line_h
        tw = text_width(label, self.ui_bold, fs)
        rect = fitz.Rect(x0, top, min(x1, x0 + pad_x + tw + 7), bottom)
        self.page.draw_rect(rect, fill=self.theme.bg_hunk, color=None, fill_opacity=0.9)
        self.page.insert_text((x0 + pad_x, self.y_base), label,
                              fontname=self.ui_bold, fontsize=fs, color=self.theme.tx_hunk)
        self.y_base = bottom + self.layout.gap_badge_to_hunk + fs

    def draw_hunk_header(self, text: str):
        assert self.page is not None and self.y_base is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        pad_x = 7

        top = self.y_base - fs
        bottom = top + line_h
        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom),
                            fill=self.theme.bg_hunk, color=None, fill_opacity=0.9)
        if text:
            self.page.insert_text((x0 + pad_x, self.y_base), text,
                                  fontname=self.mono_font, fontsize=fs, color=self.theme.tx_hunk)
        self.y_base = bottom + self.layout.gap_hunk_to_code + fs

    # -- Height measurement --------------------------------------------------

    def _gutter_sample(self) -> str:
        g = self.layout.gutter_chars
        return f"{'9' * g} {'9' * g} "

    def measure_hunk_height_unified(self, hunk: Hunk, hide_context: bool) -> float:
        assert self.page is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        gutter_w = text_width(self._gutter_sample(), self.mono_font, fs)
        max_w = max(12.0, x1 - (x0 + gutter_w + self.layout.gutter_gap))

        total = line_h + self.layout.gap_hunk_to_code   # hunk header
        for dl in hunk.lines:
            if hide_context and dl.kind == "ctx":
                continue
            total += len(_wrap_text(dl.text, max_w, self.mono_font, fs)) * line_h
        total += self.layout.block_gap_y
        return total

    def measure_file_height_unified(self, diff_file: DiffFile, hide_context: bool) -> float:
        assert self.page is not None
        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap

        gutter_w = text_width(self._gutter_sample(), self.mono_font, fs)
        max_w = max(12.0, x1 - (x0 + gutter_w + self.layout.gutter_gap))

        total = line_h + self.layout.gap_badge_to_hunk   # file badge
        for h in diff_file.hunks:
            total += line_h + self.layout.gap_hunk_to_code
            for dl in h.lines:
                if hide_context and dl.kind == "ctx":
                    continue
                total += len(_wrap_text(dl.text, max_w, self.mono_font, fs)) * line_h
            total += self.layout.block_gap_y
        total += self.layout.block_gap_y
        return total

    # -- Unified rendering ---------------------------------------------------

    def render_file_unified(self, diff_file: DiffFile, title: str, hide_context: bool, show_hunk_header: bool = True):
        self.start_if_needed(title)
        assert self.page is not None and self.y_base is not None

        # Keep-together: jump to a fresh page if the whole file fits there
        required = self.measure_file_height_unified(diff_file, hide_context)
        capacity = self.page_capacity()
        if required <= capacity and required > self.space_left():
            self.ensure_y(10_000)

        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        gutter_w = text_width(self._gutter_sample(), self.mono_font, fs)
        text_x = x0 + gutter_w + self.layout.gutter_gap

        # File badge
        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"
        self.ensure_y(line_h)
        self.draw_file_badge(label)

        for h in diff_file.hunks:
            # Keep-together for individual hunks
            h_req = self.measure_hunk_height_unified(h, hide_context)
            if h_req > self.space_left() and h_req <= self.page_capacity():
                self.ensure_y(10_000)

            self.widow_check_before_hunk(line_h)
            self.ensure_y(line_h)
            self.draw_hunk_header(h.header if show_hunk_header else "")

            for dl in h.lines:
                if hide_context and dl.kind == "ctx":
                    continue

                max_w = max(12.0, x1 - text_x)
                parts = _wrap_text(dl.text, max_w, self.mono_font, fs)
                rows_h = len(parts) * line_h

                self.ensure_y(rows_h)
                x0, _, x1, _ = self.box(self.page)
                text_x = x0 + gutter_w + self.layout.gutter_gap

                for j, part in enumerate(parts):
                    top = self.y_base - fs
                    bottom = top + line_h

                    # Background + side bar
                    if dl.kind == "add":
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom),
                                            fill=self.theme.bg_added, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(x0, top, x0 + 2.5, bottom),
                                            fill=self.theme.bar_added, color=None, fill_opacity=1.0)
                    elif dl.kind == "del":
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom),
                                            fill=self.theme.bg_removed, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(x0, top, x0 + 2.5, bottom),
                                            fill=self.theme.bar_removed, color=None, fill_opacity=1.0)
                    else:
                        self.page.draw_rect(fitz.Rect(x0, top, x1, bottom),
                                            fill=self.theme.bg_context, color=None, fill_opacity=0.9)

                    # Line numbers (first wrap-row only)
                    if j == 0:
                        g = self.layout.gutter_chars
                        old_s = f"{dl.old_num:>{g}d}" if dl.old_num is not None else " " * g
                        new_s = f"{dl.new_num:>{g}d}" if dl.new_num is not None else " " * g
                        ln_text = f"{old_s} {new_s} "
                        self.page.insert_text((x0, self.y_base), ln_text,
                                              fontname=self.mono_font, fontsize=fs,
                                              color=self.theme.ui_subtle)

                    # Code text
                    col = (self.theme.tx_added if dl.kind == "add"
                           else self.theme.tx_removed if dl.kind == "del"
                           else self.theme.tx_context)
                    self.page.insert_text((text_x, self.y_base), part,
                                          fontname=self.mono_font, fontsize=fs, color=col)
                    self.y_base += line_h

            self.y_base += self.layout.block_gap_y
        self.y_base += self.layout.block_gap_y

    # -- Side-by-side rendering ----------------------------------------------

    def render_file_sbs(self, diff_file: DiffFile, title: str, show_hunk_header: bool = True):
        self.start_if_needed(title)
        assert self.page is not None and self.y_base is not None

        x0, _, x1, _ = self.box(self.page)
        fs = self.layout.font_size
        line_h = fs + self.layout.line_gap
        gap = self.layout.col_gap
        col_w = (x1 - x0 - gap) / 2

        gutter_w = text_width(self._gutter_sample(), self.mono_font, fs)
        left_x0 = x0
        left_x1 = x0 + col_w
        right_x0 = left_x1 + gap
        right_x1 = x1
        left_text_x = left_x0 + gutter_w + self.layout.gutter_gap
        right_text_x = right_x0 + gutter_w + self.layout.gutter_gap

        # File badge
        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"
        self.ensure_y(line_h)
        self.draw_file_badge(label)

        for h in diff_file.hunks:
            self.widow_check_before_hunk(line_h)
            self.ensure_y(line_h)
            self.draw_hunk_header(h.header if show_hunk_header else "")

            i = 0
            while i < len(h.lines):
                # Pair del+add, or handle singles
                left: Optional[DiffLine] = None
                right: Optional[DiffLine] = None
                ln = h.lines[i]
                if ln.kind == "del":
                    if i + 1 < len(h.lines) and h.lines[i + 1].kind == "add":
                        left, right = ln, h.lines[i + 1]
                        i += 2
                    else:
                        left = ln
                        i += 1
                elif ln.kind == "add":
                    right = ln
                    i += 1
                else:
                    left = right = ln
                    i += 1

                l_max = max(12.0, left_x1 - left_text_x)
                r_max = max(12.0, right_x1 - right_text_x)
                l_parts = _wrap_text(left.text, l_max, self.mono_font, fs) if left else [""]
                r_parts = _wrap_text(right.text, r_max, self.mono_font, fs) if right else [""]

                rows = max(len(l_parts), len(r_parts))
                rows_h = rows * line_h
                self.ensure_y(rows_h)

                # Recalculate after possible page break
                x0, _, x1, _ = self.box(self.page)
                left_x0 = x0
                left_x1 = x0 + col_w
                right_x0 = left_x1 + gap
                right_x1 = x1
                left_text_x = left_x0 + gutter_w + self.layout.gutter_gap
                right_text_x = right_x0 + gutter_w + self.layout.gutter_gap

                for j in range(rows):
                    top = self.y_base - fs
                    bottom = top + line_h

                    # Backgrounds
                    if left and left.kind == "del":
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x1, bottom),
                                            fill=self.theme.bg_removed, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x0 + 2.5, bottom),
                                            fill=self.theme.bar_removed, color=None, fill_opacity=1.0)
                    if right and right.kind == "add":
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x1, bottom),
                                            fill=self.theme.bg_added, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x0 + 2.5, bottom),
                                            fill=self.theme.bar_added, color=None, fill_opacity=1.0)
                    if left and right and left.kind == "ctx" and right.kind == "ctx":
                        self.page.draw_rect(fitz.Rect(left_x0, top, left_x1, bottom),
                                            fill=self.theme.bg_context, color=None, fill_opacity=0.9)
                        self.page.draw_rect(fitz.Rect(right_x0, top, right_x1, bottom),
                                            fill=self.theme.bg_context, color=None, fill_opacity=0.9)

                    # Line numbers
                    if j == 0 and left and left.kind in ("ctx", "del"):
                        num = left.old_num if left.old_num is not None else 0
                        ln_txt = f"{num:>{self.layout.gutter_chars}d} "
                        self.page.insert_text((left_x0, self.y_base), ln_txt,
                                              fontname=self.mono_font, fontsize=fs, color=self.theme.ui_subtle)
                    if j == 0 and right and right.kind in ("ctx", "add"):
                        num = right.new_num if right.new_num is not None else 0
                        rn_txt = f"{num:>{self.layout.gutter_chars}d} "
                        self.page.insert_text((right_x0, self.y_base), rn_txt,
                                              fontname=self.mono_font, fontsize=fs, color=self.theme.ui_subtle)

                    # Code text
                    lp = l_parts[j] if j < len(l_parts) else ""
                    rp = r_parts[j] if j < len(r_parts) else ""
                    if left:
                        col = self.theme.tx_removed if left.kind == "del" else self.theme.tx_context
                        self.page.insert_text((left_text_x, self.y_base), lp,
                                              fontname=self.mono_font, fontsize=fs, color=col)
                    if right:
                        col = self.theme.tx_added if right.kind == "add" else self.theme.tx_context
                        self.page.insert_text((right_text_x, self.y_base), rp,
                                              fontname=self.mono_font, fontsize=fs, color=col)
                    self.y_base += line_h

            self.y_base += self.layout.block_gap_y
        self.y_base += self.layout.block_gap_y

    # -- Save ----------------------------------------------------------------

    def save(self, output_path: str):
        self.draw_footer_page_numbers()
        self.doc.save(output_path)
        self.doc.close()


# ---------------------------------------------------------------------------
#  Word Renderer
# ---------------------------------------------------------------------------

def render_word(
    all_files:    List[DiffFile],
    title:        str,
    output_path:  str,
    theme:        Theme,
    hide_context: bool,
    view:         str,
    show_hunk_header: bool = True,
    font_size_pt: float = 9.5,
):
    """Render parsed diff files into a Word .docx document (requires python-docx)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[ERROR] python-docx is not installed. Run: pip install python-docx",
              file=sys.stderr)
        sys.exit(1)

    # -- Word helper functions -----------------------------------------------

    def hex_color(rgb_tuple: Tuple[float, float, float]) -> str:
        return _rgb_to_hex(rgb_tuple)

    def set_cell_bg(cell, hex_str: str):
        """Set table cell background colour via XML shading."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_str)
        tcPr.append(shd)

    def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
        """Set cell margins in twips (1 pt = 20 twips)."""
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    def set_cell_width(cell, twips: int):
        """Set exact cell width in twips."""
        tcPr = cell._element.get_or_add_tcPr()
        old_tcW = tcPr.find(qn("w:tcW"))
        if old_tcW is not None:
            tcPr.remove(old_tcW)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(twips))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    def set_table_full_width(table):
        """Set table to 100% page width."""
        tbl_el = table._tbl
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tblPr)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")
        tblPr.append(tblW)

    def setup_fixed_2col_table(tbl, col0_twips: int, col1_twips: int):
        """Configure a 2-column table with fixed layout, full width, and exact column widths."""
        tbl_el = tbl._tbl
        tblPr = tbl_el.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_el.insert(0, tblPr)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "fixed")
        tblPr.append(tblLayout)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")
        tblPr.append(tblW)
        tblCS = OxmlElement("w:tblCellSpacing")
        tblCS.set(qn("w:w"), "0")
        tblCS.set(qn("w:type"), "dxa")
        tblPr.append(tblCS)
        tblGrid = tbl_el.find(qn("w:tblGrid"))
        if tblGrid is not None:
            tbl_el.remove(tblGrid)
        tblGrid = OxmlElement("w:tblGrid")
        gc0 = OxmlElement("w:gridCol")
        gc0.set(qn("w:w"), str(col0_twips))
        gc1 = OxmlElement("w:gridCol")
        gc1.set(qn("w:w"), str(col1_twips))
        tblGrid.append(gc0)
        tblGrid.append(gc1)
        tbl_el.insert(1, tblGrid)

    def remove_table_borders(table):
        """Remove all borders from a table (table-level + cell-level)."""
        border_names = ("top", "left", "bottom", "right", "insideH", "insideV")

        # Table-level
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for name in border_names:
            el = OxmlElement(f"w:{name}")
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
            tblBorders.append(el)

        # Cell-level
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn("w:tcBorders"))
                if tcBorders is None:
                    tcBorders = OxmlElement("w:tcBorders")
                    tcPr.append(tcBorders)
                for name in border_names:
                    el = OxmlElement(f"w:{name}")
                    el.set(qn("w:val"), "none")
                    el.set(qn("w:sz"), "0")
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "auto")
                    tcBorders.append(el)

    def add_colored_paragraph(doc, text: str, bg_hex: str, fg_rgb,
                              bold: bool = False, mono: bool = True,
                              indent_pt: float = 0):
        """Add a single-row table acting as a highlighted paragraph."""
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        remove_table_borders(table)
        set_table_full_width(table)

        cell = table.cell(0, 0)
        set_cell_bg(cell, bg_hex)
        set_cell_margins(cell, left=40)

        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if indent_pt:
            para.paragraph_format.left_indent = Pt(indent_pt)
        run = para.add_run(text)
        run.font.name = "Consolas" if mono else "Calibri"
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(
            int(fg_rgb[0] * 255), int(fg_rgb[1] * 255), int(fg_rgb[2] * 255))
        return table

    def fill_sbs_cell(cell, dl: Optional[DiffLine], side: str):
        """Fill one side of a side-by-side row."""
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if dl is None:
            return
        num = dl.old_num if side == "left" else dl.new_num
        num_str = f"{num:>5d}  " if num is not None else "        "

        run_num = para.add_run(num_str)
        run_num.font.name = "Consolas"
        run_num.font.size = Pt(font_size_pt)
        run_num.font.color.rgb = RGBColor(125, 125, 125)

        fg = (theme.tx_removed if dl.kind == "del"
              else theme.tx_added if dl.kind == "add"
              else theme.tx_context)
        run_txt = para.add_run(dl.text)
        run_txt.font.name = "Consolas"
        run_txt.font.size = Pt(font_size_pt)
        run_txt.font.color.rgb = RGBColor(
            int(fg[0] * 255), int(fg[1] * 255), int(fg[2] * 255))

    # -- Build the document ---------------------------------------------------

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin   = Inches(0.7)
        section.right_margin  = Inches(0.7)

    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(40, 40, 40)

    stamp_para = doc.add_paragraph(dt.datetime.now().strftime("Generated: %Y-%m-%d %H:%M"))
    stamp_para.runs[0].font.size = Pt(font_size_pt - 1)
    stamp_para.runs[0].font.color.rgb = RGBColor(125, 125, 125)
    doc.add_paragraph("")  # spacer

    bg_add_hex  = hex_color(theme.bg_added)
    bg_del_hex  = hex_color(theme.bg_removed)
    bg_ctx_hex  = hex_color(theme.bg_context)
    bg_hunk_hex = hex_color(theme.bg_hunk)

    # Padding between line-number gutter and code column (in twips, 1 pt = 20 twips).
    gutter_right_pad = 50
    code_left_pad = 60

    # Pre-compute gutter width in twips for unified 2-col tables.
    # Gutter: 9 Consolas chars (4+1+4) at ~60% of font-size points, in twips.
    gutter_twips = int(font_size_pt * 0.6 * 9 * 20) + 40 + gutter_right_pad

    # -- Per-file rendering ---------------------------------------------------

    for diff_file in all_files:
        label = diff_file.new_path or diff_file.old_path or "(Unnamed)"

        # File badge
        add_colored_paragraph(doc, label, bg_hunk_hex, theme.tx_hunk, bold=True, mono=False)

        for h in diff_file.hunks:
            # Hunk separator
            hunk_label = h.header if show_hunk_header else ""
            add_colored_paragraph(doc, hunk_label, bg_hunk_hex, theme.tx_hunk, bold=False, mono=True)

            # Side-by-side view
            if view == "side-by-side":
                i = 0
                while i < len(h.lines):
                    left: Optional[DiffLine] = None
                    right: Optional[DiffLine] = None
                    ln = h.lines[i]
                    if ln.kind == "del":
                        if i + 1 < len(h.lines) and h.lines[i + 1].kind == "add":
                            left, right = ln, h.lines[i + 1]
                            i += 2
                        else:
                            left = ln
                            i += 1
                    elif ln.kind == "add":
                        right = ln
                        i += 1
                    else:
                        left = right = ln
                        i += 1

                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.style = "Table Grid"
                    remove_table_borders(tbl)
                    l_cell, r_cell = tbl.cell(0, 0), tbl.cell(0, 1)

                    set_cell_bg(l_cell, bg_del_hex if left and left.kind == "del" else bg_ctx_hex)
                    set_cell_bg(r_cell, bg_add_hex if right and right.kind == "add" else bg_ctx_hex)
                    set_cell_margins(l_cell, left=30)
                    set_cell_margins(r_cell, left=30)

                    fill_sbs_cell(l_cell, left, "left")
                    fill_sbs_cell(r_cell, right, "right")

            # Unified view
            else:
                for dl in h.lines:
                    if hide_context and dl.kind == "ctx":
                        continue

                    bg_hex = (bg_add_hex if dl.kind == "add"
                              else bg_del_hex if dl.kind == "del"
                              else bg_ctx_hex)
                    fg = (theme.tx_added if dl.kind == "add"
                          else theme.tx_removed if dl.kind == "del"
                          else theme.tx_context)

                    old_s = f"{dl.old_num:>4d}" if dl.old_num is not None else "    "
                    new_s = f"{dl.new_num:>4d}" if dl.new_num is not None else "    "

                    tbl = doc.add_table(rows=1, cols=2)
                    tbl.style = "Table Grid"
                    remove_table_borders(tbl)
                    code_twips = 9892 - gutter_twips
                    setup_fixed_2col_table(tbl, gutter_twips, code_twips)

                    c_num = tbl.cell(0, 0)
                    c_code = tbl.cell(0, 1)

                    set_cell_bg(c_num, bg_hex)
                    set_cell_bg(c_code, bg_hex)
                    set_cell_width(c_num, gutter_twips)
                    set_cell_width(c_code, code_twips)
                    set_cell_margins(c_num, right=gutter_right_pad)
                    set_cell_margins(c_code, left=code_left_pad)

                    # Line numbers
                    p_num = c_num.paragraphs[0]
                    p_num.paragraph_format.space_before = Pt(0)
                    p_num.paragraph_format.space_after = Pt(0)
                    r_num = p_num.add_run(f"{old_s} {new_s}")
                    r_num.font.name = "Consolas"
                    r_num.font.size = Pt(font_size_pt)
                    r_num.font.color.rgb = RGBColor(125, 125, 125)

                    # Code text (wraps naturally within its cell)
                    p_code = c_code.paragraphs[0]
                    p_code.paragraph_format.space_before = Pt(0)
                    p_code.paragraph_format.space_after = Pt(0)

                    # Strip leading whitespace → use paragraph indent instead,
                    # so Word continuation lines align with the code start.
                    code_text = dl.text
                    leading = len(code_text) - len(code_text.lstrip(" "))
                    if leading > 0:
                        # ~60% of font-size per monospace char, in twips (1 pt = 20 twips)
                        indent_tw = int(leading * font_size_pt * 0.6 * 20)
                        pPr_code = p_code._element.get_or_add_pPr()
                        ind = OxmlElement("w:ind")
                        ind.set(qn("w:left"), str(indent_tw))
                        pPr_code.append(ind)
                        code_text = code_text.lstrip(" ")

                    r_code = p_code.add_run(code_text)
                    r_code.font.name = "Consolas"
                    r_code.font.size = Pt(font_size_pt)
                    r_code.font.color.rgb = RGBColor(
                        int(fg[0] * 255), int(fg[1] * 255), int(fg[2] * 255))

        doc.add_paragraph("")  # spacer between files

    doc.save(output_path)
    print(f"✓ Word document created: {output_path}")

# ---------------------------------------------------------------------------
#  CLI
# ---------------------------------------------------------------------------

def _build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="PR-style PDF (and optionally Word .docx) from unified git diffs.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    # Positional
    p.add_argument("inputs", nargs="+",
                   help="Diff file(s) or '-' for STDIN")
    # Output
    p.add_argument("-o", "--output", required=True,
                   help="Output PDF path (e.g. diff.pdf)")
    p.add_argument("--word", action="store_true",
                   help="Also generate a Word .docx alongside the PDF")
    p.add_argument("--word-output", default=None, metavar="FILE.docx",
                   help="Custom path for the .docx (default: same base as --output)")
    # Layout
    p.add_argument("--title", default="Changed Code",
                   help="Document title")
    p.add_argument("--view", choices=["unified", "side-by-side"], default="unified",
                   help="Layout mode")
    p.add_argument("--hide-context", action="store_true",
                   help="Hide context lines (show only +/-) in unified view")
    p.add_argument("--hunk-header", choices=["show", "hide"], default="show",
                   help="Show or hide @@ hunk headers (hide = empty blue separator)")
    p.add_argument("--landscape", action="store_true",
                   help="A4 landscape")
    p.add_argument("--font-size", type=float, default=9.5,
                   help="Monospace font size (pt)")
    p.add_argument("--tabsize", type=int, default=4,
                   help="Tab expansion width")
    p.add_argument("--theme", choices=["light", "dark"], default="light",
                   help="Colour theme")
    # Font overrides
    p.add_argument("--mono-font-file", default=None,
                   help="TTF/OTF path for monospace font")
    p.add_argument("--mono-bold-font-file", default=None,
                   help="TTF/OTF path for monospace bold font")
    p.add_argument("--ui-font-file", default=None,
                   help="TTF/OTF path for UI font")
    p.add_argument("--ui-bold-font-file", default=None,
                   help="TTF/OTF path for UI bold font")
    # Debug
    p.add_argument("--debug", action="store_true",
                   help="Print parser debug output to stderr")
    return p

# ---------------------------------------------------------------------------
#  Main
# ---------------------------------------------------------------------------

def main():
    args = _build_parser().parse_args()
    theme = LIGHT if args.theme == "light" else DARK

    # Fonts: user overrides > system detection > Courier fallback
    sys_fonts = detect_system_fonts()
    fonts = Fonts(
        ui        = _load_font(args.ui_font_file)        or sys_fonts.ui,
        ui_bold   = _load_font(args.ui_bold_font_file)   or sys_fonts.ui_bold,
        mono      = _load_font(args.mono_font_file)      or sys_fonts.mono,
        mono_bold = _load_font(args.mono_bold_font_file)  or sys_fonts.mono_bold,
    )

    # Parse diff inputs
    all_files: List[DiffFile] = []
    for path in args.inputs:
        try:
            txt = read_text(path)
        except FileNotFoundError:
            print(f"[WARN] File not found: {path}", file=sys.stderr)
            continue
        all_files.extend(parse_unified_diff(txt, tabsize=args.tabsize, debug=args.debug))

    if not all_files:
        print("[ERROR] No parsable diffs found.", file=sys.stderr)
        print("Hints:", file=sys.stderr)
        print("  • Use a unified diff: e.g. `git diff <commit>` or `git show <commit>`", file=sys.stderr)
        print("  • Not supported: --word-diff, --name-only, --name-status", file=sys.stderr)
        sys.exit(2)

    # Render PDF
    layout = Layout(font_size=args.font_size)
    renderer = Renderer(theme=theme, layout=layout, landscape=args.landscape, fonts=fonts)
    show_hunk = args.hunk_header == "show"
    for df in all_files:
        if args.view == "unified":
            renderer.render_file_unified(df, title=args.title, hide_context=args.hide_context, show_hunk_header=show_hunk)
        else:
            renderer.render_file_sbs(df, title=args.title, show_hunk_header=show_hunk)
    renderer.save(args.output)

    # Render Word (optional)
    if args.word or args.word_output:
        word_path = args.word_output
        if not word_path:
            base, _ = os.path.splitext(args.output)
            word_path = base + ".docx"
        render_word(
            all_files=all_files,
            title=args.title,
            output_path=word_path,
            theme=theme,
            hide_context=args.hide_context,
            view=args.view,
            show_hunk_header=show_hunk,
            font_size_pt=args.font_size,
        )


if __name__ == "__main__":
    main()
